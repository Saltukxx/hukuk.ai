"""
Document Generator Service
This module handles the generation of legal documents from templates.
"""

import os
import re
import json
import uuid
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging

logger = logging.getLogger(__name__)

class DocumentGenerator:
    """Handles the generation of legal documents from templates"""

    def __init__(self):
        """Initialize the document generator service"""
        self.output_dir = os.path.join("app", "output")
        # Ensure output directory exists
        os.makedirs(self.output_dir, exist_ok=True)
        logger.info("DocumentGenerator initialized")

    async def generate_document(self, template_name, template_data, output_format="docx"):
        """
        Generate a document based on a template and provided data.
        
        Args:
            template_name (str): The name of the template to use
            template_data (dict): The data to fill the template with
            output_format (str): The output format (currently only 'docx' is supported)
            
        Returns:
            dict: Results including the document path
        """
        logger.info(f"Generating document for template: {template_name}")
        
        # Sanitize inputs to prevent path traversal
        template_name = self._sanitize_filename(template_name)
        
        # Generate a unique document ID
        document_id = str(uuid.uuid4())
        
        # Generate the document
        docx_file_path = await self._generate_docx(template_name, template_data, document_id)
        
        # Return the result
        return {
            "document_id": document_id,
            "document_path": docx_file_path,
            "template_name": template_name,
            "created_at": datetime.now().isoformat()
        }

    async def _generate_docx(self, template_name, template_data, document_id):
        """
        Generate a Microsoft Word document from template data
        
        Args:
            template_name (str): The name of the template to use
            template_data (dict): The data to fill the template with
            document_id (str): The unique document ID
            
        Returns:
            str: The path to the generated document
        """
        # Create a new Word document
        doc = Document()
        
        # Set margin
        for section in doc.sections:
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
        
        if template_name == "dilekce":
            return await self._generate_dilekce(doc, template_data, document_id)
        elif template_name == "ihtarname":
            return await self._generate_ihtarname(doc, template_data, document_id)
        elif template_name == "vekaletname":
            return await self._generate_vekaletname(doc, template_data, document_id)
        elif template_name == "dava_dilekce":
            return await self._generate_dava_dilekce(doc, template_data, document_id)
        else:
            return await self._generate_generic(doc, template_name, template_data, document_id)
    
    async def _generate_dilekce(self, doc, template_data, document_id):
        """Generate a petition document"""
        # Title
        title = doc.add_paragraph()
        title_run = title.add_run("DİLEKÇE")
        title_run.bold = True
        title_run.font.size = Pt(14)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Recipient
        recipient = doc.add_paragraph()
        recipient_run = recipient.add_run(template_data.get("kurum", "").upper())
        recipient_run.bold = True
        recipient.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Subject
        subject = doc.add_paragraph()
        subject.add_run("Konu: ").bold = True
        subject.add_run(template_data.get("konu", ""))
        
        # Content
        doc.add_paragraph().add_run()  # Empty line
        content = doc.add_paragraph()
        content.add_run(template_data.get("icerik", ""))
        
        # Signature
        doc.add_paragraph().add_run()  # Empty line
        signature = doc.add_paragraph()
        signature.add_run("Saygılarımla,")
        signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Name
        name = doc.add_paragraph()
        name.add_run(template_data.get("ad_soyad", ""))
        name.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Attachments
        if "ekler" in template_data and template_data["ekler"]:
            doc.add_paragraph().add_run()  # Empty line
            attachments = doc.add_paragraph()
            attachments.add_run("Ekler:").bold = True
            if isinstance(template_data["ekler"], list):
                for i, attachment in enumerate(template_data["ekler"], 1):
                    attachments.add_run(f"\n{i}. {attachment}")
            else:
                attachments.add_run(f"\n{template_data['ekler']}")
        
        # Save the document
        file_path = os.path.join(self.output_dir, f"dilekce_{document_id}.docx")
        doc.save(file_path)
        
        return file_path
    
    async def _generate_ihtarname(self, doc, template_data, document_id):
        """Generate a formal warning document"""
        # Title
        title = doc.add_paragraph()
        title_run = title.add_run("İHTARNAME")
        title_run.bold = True
        title_run.font.size = Pt(14)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Sender
        sender_title = doc.add_paragraph()
        sender_title.add_run("Gönderen: ").bold = True
        sender = doc.add_paragraph()
        sender.add_run(template_data.get("gonderen", ""))
        
        # Recipient
        recipient_title = doc.add_paragraph()
        recipient_title.add_run("Muhatap: ").bold = True
        recipient = doc.add_paragraph()
        recipient.add_run(template_data.get("alici", ""))
        
        # Subject
        subject = doc.add_paragraph()
        subject.add_run("Konu: ").bold = True
        subject.add_run(template_data.get("konu", ""))
        
        # Content
        doc.add_paragraph().add_run()  # Empty line
        content = doc.add_paragraph()
        content.add_run(template_data.get("icerik", ""))
        
        # Conclusion
        doc.add_paragraph().add_run()  # Empty line
        conclusion = doc.add_paragraph()
        conclusion.add_run("Sonuç ve Talep: ").bold = True
        conclusion.add_run(template_data.get("sonuc_talep", ""))
        
        # Signature
        doc.add_paragraph().add_run()  # Empty line
        signature = doc.add_paragraph()
        signature.add_run("Saygılarımla,")
        signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Name
        name = doc.add_paragraph()
        name.add_run(template_data.get("ad_soyad", ""))
        name.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Save the document
        file_path = os.path.join(self.output_dir, f"ihtarname_{document_id}.docx")
        doc.save(file_path)
        
        return file_path
    
    async def _generate_vekaletname(self, doc, template_data, document_id):
        """Generate a power of attorney document"""
        # Title
        title = doc.add_paragraph()
        title_run = title.add_run("VEKALETNAME")
        title_run.bold = True
        title_run.font.size = Pt(14)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Principal
        principal_title = doc.add_paragraph()
        principal_title.add_run("Vekil Eden: ").bold = True
        principal = doc.add_paragraph()
        principal.add_run(template_data.get("vekil_eden", ""))
        
        # Attorney
        attorney_title = doc.add_paragraph()
        attorney_title.add_run("Vekil: ").bold = True
        attorney = doc.add_paragraph()
        attorney.add_run(template_data.get("vekil", ""))
        
        # Content
        doc.add_paragraph().add_run()  # Empty line
        content = doc.add_paragraph()
        content.add_run(template_data.get("icerik", ""))
        
        # Powers
        doc.add_paragraph().add_run()  # Empty line
        powers_title = doc.add_paragraph()
        powers_title.add_run("Verilen Yetkiler:").bold = True
        
        if "yetkiler" in template_data:
            powers = doc.add_paragraph()
            if isinstance(template_data["yetkiler"], list):
                for i, power in enumerate(template_data["yetkiler"], 1):
                    powers.add_run(f"{i}. {power}\n")
            else:
                powers.add_run(template_data["yetkiler"])
        
        # Signature
        doc.add_paragraph().add_run()  # Empty line
        signature = doc.add_paragraph()
        signature.add_run("Tarih: " + datetime.now().strftime("%d/%m/%Y"))
        signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Name
        name = doc.add_paragraph()
        name.add_run(template_data.get("ad_soyad", ""))
        name.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Save the document
        file_path = os.path.join(self.output_dir, f"vekaletname_{document_id}.docx")
        doc.save(file_path)
        
        return file_path
    
    async def _generate_dava_dilekce(self, doc, template_data, document_id):
        """Generate a lawsuit petition document"""
        # Title
        title = doc.add_paragraph()
        title_run = title.add_run(template_data.get("mahkeme", "").upper())
        title_run.bold = True
        title_run.font.size = Pt(12)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph()
        subtitle_run = subtitle.add_run(template_data.get("dava_turu", "").upper() + " DAVASI DİLEKÇESİ")
        subtitle_run.bold = True
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Parties
        doc.add_paragraph().add_run()  # Empty line
        plaintiff_title = doc.add_paragraph()
        plaintiff_title.add_run("Davacı: ").bold = True
        plaintiff = doc.add_paragraph()
        plaintiff.add_run(template_data.get("davaci", ""))
        
        defendant_title = doc.add_paragraph()
        defendant_title.add_run("Davalı: ").bold = True
        defendant = doc.add_paragraph()
        defendant.add_run(template_data.get("davali", ""))
        
        # Subject
        subject = doc.add_paragraph()
        subject.add_run("Konu: ").bold = True
        subject.add_run(template_data.get("konu", ""))
        
        # Value
        value = doc.add_paragraph()
        value.add_run("Dava Değeri: ").bold = True
        value.add_run(template_data.get("deger", ""))
        
        # Content
        doc.add_paragraph().add_run()  # Empty line
        content = doc.add_paragraph()
        content.add_run(template_data.get("aciklamalar", ""))
        
        # Evidence
        if "deliller" in template_data:
            doc.add_paragraph().add_run()  # Empty line
            evidence_title = doc.add_paragraph()
            evidence_title.add_run("Deliller:").bold = True
            evidence = doc.add_paragraph()
            if isinstance(template_data["deliller"], list):
                for i, item in enumerate(template_data["deliller"], 1):
                    evidence.add_run(f"{i}. {item}\n")
            else:
                evidence.add_run(template_data["deliller"])
        
        # Legal basis
        if "hukuki_sebepler" in template_data:
            doc.add_paragraph().add_run()  # Empty line
            legal_title = doc.add_paragraph()
            legal_title.add_run("Hukuki Sebepler:").bold = True
            legal = doc.add_paragraph()
            legal.add_run(template_data.get("hukuki_sebepler", ""))
        
        # Request
        doc.add_paragraph().add_run()  # Empty line
        request_title = doc.add_paragraph()
        request_title.add_run("Sonuç ve Talep:").bold = True
        request = doc.add_paragraph()
        request.add_run(template_data.get("talep", ""))
        
        # Signature
        doc.add_paragraph().add_run()  # Empty line
        signature = doc.add_paragraph()
        signature.add_run("Saygılarımla,")
        signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Name
        name = doc.add_paragraph()
        name.add_run(template_data.get("ad_soyad", ""))
        name.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Save the document
        file_path = os.path.join(self.output_dir, f"dava_dilekce_{document_id}.docx")
        doc.save(file_path)
        
        return file_path
    
    async def _generate_generic(self, doc, template_name, template_data, document_id):
        """Generate a generic document based on template data"""
        # Title
        title = doc.add_paragraph()
        title_run = title.add_run(template_name.upper().replace("_", " "))
        title_run.bold = True
        title_run.font.size = Pt(14)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add all template data as content
        doc.add_paragraph().add_run()  # Empty line
        
        for key, value in template_data.items():
            if value:
                field = doc.add_paragraph()
                field.add_run(key.replace("_", " ").title() + ": ").bold = True
                
                if isinstance(value, list):
                    for i, item in enumerate(value, 1):
                        field.add_run(f"\n{i}. {item}")
                else:
                    field.add_run(str(value))
        
        # Save the document
        file_path = os.path.join(self.output_dir, f"{template_name}_{document_id}.docx")
        doc.save(file_path)
        
        return file_path
    
    def _sanitize_filename(self, filename):
        """
        Sanitize a filename to prevent path traversal attacks
        """
        # Remove any directory components
        filename = os.path.basename(filename)
        
        # Remove any potentially dangerous characters
        filename = re.sub(r'[^\w\-_.]', '_', filename)
        
        return filename 
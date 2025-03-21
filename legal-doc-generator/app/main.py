from fastapi import FastAPI, HTTPException, BackgroundTasks, Query
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from fastapi import Request
from pydantic import BaseModel
from typing import Dict, Any, Optional, List
import uuid
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.services.ai_service import AIService
from app.services.legal_database_service import LegalDatabaseService
from app.services.data_collector_service import LegalDataCollector
import re
import sqlite3

app = FastAPI(
    title="Hukuk.AI - Türkçe Hukuki Belge Üretici",
    description="Dilekçe ve hukuki belge hazırlama sistemi",
    version="1.0.0"
)

# CORS ayarları
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Tüm originlere izin ver (geliştirme için)
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Statik dosyaları sunmak için
app.mount("/static", StaticFiles(directory="app/static"), name="static")

# Templates
templates = Jinja2Templates(directory="app/templates")

# Hizmetlerin başlatılması
ai_service = AIService()
legal_db_service = LegalDatabaseService()
data_collector_service = LegalDataCollector()

class GenerateDocumentRequest(BaseModel):
    template_name: str
    template_data: Dict[str, Any]
    metadata: Optional[Dict[str, Any]] = None

class GenerateAIEnhancedDocumentRequest(BaseModel):
    template_name: str
    template_data: Dict[str, Any]
    case_description: str
    case_category: str
    metadata: Optional[Dict[str, Any]] = None

class GenerateResponse(BaseModel):
    document_id: str
    download_url: str
    metadata: Dict[str, Any]

class TemplateInfo(BaseModel):
    id: str
    name: str
    description: str
    category: str
    required_fields: List[Dict[str, Any]]

class TemplateListResponse(BaseModel):
    templates: List[TemplateInfo]

class CaseAnalysisRequest(BaseModel):
    case_description: str
    case_category: str

class CaseAnalysisResponse(BaseModel):
    analysis: Dict[str, Any]
    relevant_laws: List[Dict[str, Any]]
    relevant_decisions: List[Dict[str, Any]]

class DocumentGenerator:
    def generate_document(self, template_name: str, template_data: Dict[str, Any]) -> bytes:
        # Template tipine göre doğru generator metodunu çağır
        if template_name == "dilekce":
            return self.generate_dilekce(template_data)
        elif template_name == "ihtarname":
            return self.generate_ihtarname(template_data)
        elif template_name == "vekaletname":
            return self.generate_vekaletname(template_data)
        elif template_name == "dava_dilekce":
            return self.generate_dava_dilekce(template_data)
        elif template_name == "temyiz_dilekce":
            return self.generate_temyiz_dilekce(template_data)
        else:
            raise ValueError(f"Bilinmeyen şablon: {template_name}")

    def generate_dilekce(self, template_data: Dict[str, Any]) -> bytes:
        doc = Document()
        
        # Sayfa kenar boşlukları (cm)
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # Başlık
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        title.add_run(template_data.get("kurum", "")).bold = True

        # Boşluk
        doc.add_paragraph()

        # Konu başlığı
        konu = doc.add_paragraph()
        konu.alignment = WD_ALIGN_PARAGRAPH.CENTER
        konu.add_run("KONU: ").bold = True
        konu.add_run(template_data.get("konu", ""))

        # Boşluk
        doc.add_paragraph()

        # İçerik
        content = doc.add_paragraph()
        content.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        content.add_run(template_data.get("icerik", ""))

        # Tarih ve imza
        doc.add_paragraph()
        signature = doc.add_paragraph()
        signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date = datetime.now().strftime("%d/%m/%Y")
        signature.add_run(f"Tarih: {date}\n\n")
        
        # Ad Soyad ve İmza
        signature.add_run(f"Ad Soyad: {template_data.get('ad_soyad', '')}\n\n")
        signature.add_run("İmza: _________________")

        # Ekler bölümü
        if "ekler" in template_data and template_data["ekler"]:
            doc.add_paragraph()
            ekler = doc.add_paragraph()
            ekler.add_run("EKLER:\n").bold = True
            for i, ek in enumerate(template_data["ekler"], 1):
                ekler.add_run(f"{i}- {ek}\n")

        # Belgeyi byte olarak kaydet
        import io
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()

    def generate_ihtarname(self, template_data: Dict[str, Any]) -> bytes:
        doc = Document()
        
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # Başlık
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.add_run("İHTARNAME").bold = True
        
        # Gönderen bilgileri
        sender = doc.add_paragraph()
        sender.alignment = WD_ALIGN_PARAGRAPH.LEFT
        sender.add_run("İHTAR EDEN: ").bold = True
        sender.add_run(template_data.get("gonderen", ""))
        
        # Alıcı bilgileri
        receiver = doc.add_paragraph()
        receiver.alignment = WD_ALIGN_PARAGRAPH.LEFT
        receiver.add_run("MUHATAP: ").bold = True
        receiver.add_run(template_data.get("alici", ""))
        
        # Konu
        konu = doc.add_paragraph()
        konu.alignment = WD_ALIGN_PARAGRAPH.LEFT
        konu.add_run("KONU: ").bold = True
        konu.add_run(template_data.get("konu", ""))
        
        # İçerik
        doc.add_paragraph()
        content = doc.add_paragraph()
        content.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        content.add_run(template_data.get("icerik", ""))
        
        # Sonuç ve talep
        doc.add_paragraph()
        sonuc = doc.add_paragraph()
        sonuc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        sonuc.add_run("SONUÇ VE TALEP: ").bold = True
        sonuc.add_run(template_data.get("sonuc_talep", ""))
        
        # Tarih ve imza
        doc.add_paragraph()
        signature = doc.add_paragraph()
        signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date = datetime.now().strftime("%d/%m/%Y")
        signature.add_run(f"Tarih: {date}\n\n")
        signature.add_run(f"Ad Soyad: {template_data.get('ad_soyad', '')}\n\n")
        signature.add_run("İmza: _________________")
        
        # Belgeyi byte olarak kaydet
        import io
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()
    
    def generate_vekaletname(self, template_data: Dict[str, Any]) -> bytes:
        doc = Document()
        
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # Başlık
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.add_run("VEKALETNAME").bold = True
        
        doc.add_paragraph()
        
        # Vekil Eden bilgileri
        vekil_eden = doc.add_paragraph()
        vekil_eden.alignment = WD_ALIGN_PARAGRAPH.LEFT
        vekil_eden.add_run("VEKİL EDEN: ").bold = True
        vekil_eden.add_run(template_data.get("vekil_eden", ""))
        
        # Vekil bilgileri
        vekil = doc.add_paragraph()
        vekil.alignment = WD_ALIGN_PARAGRAPH.LEFT
        vekil.add_run("VEKİL: ").bold = True
        vekil.add_run(template_data.get("vekil", ""))
        
        # Vekaletname içeriği
        doc.add_paragraph()
        content = doc.add_paragraph()
        content.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        content.add_run(template_data.get("icerik", ""))
        
        # Yetkiler
        doc.add_paragraph()
        yetkiler = doc.add_paragraph()
        yetkiler.alignment = WD_ALIGN_PARAGRAPH.LEFT
        yetkiler.add_run("YETKİLER: ").bold = True
        
        if "yetkiler" in template_data and template_data["yetkiler"]:
            for i, yetki in enumerate(template_data["yetkiler"], 1):
                yetki_paragraph = doc.add_paragraph()
                yetki_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                yetki_paragraph.paragraph_format.left_indent = Cm(1)
                yetki_paragraph.add_run(f"{i}- {yetki}")
        
        # Tarih ve imza
        doc.add_paragraph()
        signature = doc.add_paragraph()
        signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date = datetime.now().strftime("%d/%m/%Y")
        signature.add_run(f"Tarih: {date}\n\n")
        signature.add_run(f"Ad Soyad: {template_data.get('ad_soyad', '')}\n\n")
        signature.add_run("İmza: _________________")
        
        # Belgeyi byte olarak kaydet
        import io
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()
        
    def generate_dava_dilekce(self, template_data: Dict[str, Any]) -> bytes:
        doc = Document()
        
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # Başlık
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        title.add_run(template_data.get("mahkeme", "")).bold = True
        
        # Boşluk
        doc.add_paragraph()
        
        # Dava türü
        dava_turu = doc.add_paragraph()
        dava_turu.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dava_turu.add_run(template_data.get("dava_turu", "")).bold = True
        
        # Boşluk
        doc.add_paragraph()
        
        # Davacı bilgileri
        davaci = doc.add_paragraph()
        davaci.alignment = WD_ALIGN_PARAGRAPH.LEFT
        davaci.add_run("DAVACI: ").bold = True
        davaci.add_run(template_data.get("davaci", ""))
        
        # Davalı bilgileri
        davali = doc.add_paragraph()
        davali.alignment = WD_ALIGN_PARAGRAPH.LEFT
        davali.add_run("DAVALI: ").bold = True
        davali.add_run(template_data.get("davali", ""))
        
        # Konu
        konu = doc.add_paragraph()
        konu.alignment = WD_ALIGN_PARAGRAPH.LEFT
        konu.add_run("KONU: ").bold = True
        konu.add_run(template_data.get("konu", ""))
        
        # Değer
        deger = doc.add_paragraph()
        deger.alignment = WD_ALIGN_PARAGRAPH.LEFT
        deger.add_run("DEĞER: ").bold = True
        deger.add_run(template_data.get("deger", ""))
        
        # Açıklamalar
        doc.add_paragraph()
        aciklamalar_baslik = doc.add_paragraph()
        aciklamalar_baslik.alignment = WD_ALIGN_PARAGRAPH.CENTER
        aciklamalar_baslik.add_run("AÇIKLAMALAR").bold = True
        
        aciklamalar = doc.add_paragraph()
        aciklamalar.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        aciklamalar.add_run(template_data.get("aciklamalar", ""))
        
        # Hukuki sebepler
        doc.add_paragraph()
        hukuki_sebepler = doc.add_paragraph()
        hukuki_sebepler.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hukuki_sebepler.add_run("HUKUKİ SEBEPLER: ").bold = True
        hukuki_sebepler.add_run(template_data.get("hukuki_sebepler", ""))
        
        # Deliller
        doc.add_paragraph()
        deliller = doc.add_paragraph()
        deliller.alignment = WD_ALIGN_PARAGRAPH.LEFT
        deliller.add_run("DELİLLER: ").bold = True
        
        if "deliller" in template_data and template_data["deliller"]:
            for i, delil in enumerate(template_data["deliller"], 1):
                delil_paragraph = doc.add_paragraph()
                delil_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                delil_paragraph.paragraph_format.left_indent = Cm(1)
                delil_paragraph.add_run(f"{i}- {delil}")
        
        # Sonuç ve Talep
        doc.add_paragraph()
        sonuc_talep = doc.add_paragraph()
        sonuc_talep.alignment = WD_ALIGN_PARAGRAPH.LEFT
        sonuc_talep.add_run("SONUÇ VE TALEP: ").bold = True
        sonuc_talep.add_run(template_data.get("sonuc_talep", ""))
        
        # Tarih ve imza
        doc.add_paragraph()
        signature = doc.add_paragraph()
        signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date = datetime.now().strftime("%d/%m/%Y")
        signature.add_run(f"Tarih: {date}\n\n")
        signature.add_run(f"Davacı: {template_data.get('davaci_ad_soyad', '')}\n\n")
        signature.add_run("İmza: _________________")
        
        # Ekler
        if "ekler" in template_data and template_data["ekler"]:
            doc.add_paragraph()
            ekler = doc.add_paragraph()
            ekler.add_run("EKLER:\n").bold = True
            for i, ek in enumerate(template_data["ekler"], 1):
                ekler.add_run(f"{i}- {ek}\n")
        
        # Belgeyi byte olarak kaydet
        import io
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()
        
    def generate_temyiz_dilekce(self, template_data: Dict[str, Any]) -> bytes:
        doc = Document()
        
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # Başlık
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        title.add_run(template_data.get("mahkeme", "")).bold = True
        
        # Boşluk
        doc.add_paragraph()
        
        # Temyiz başlığı
        temyiz_baslik = doc.add_paragraph()
        temyiz_baslik.alignment = WD_ALIGN_PARAGRAPH.CENTER
        temyiz_baslik.add_run("TEMYİZ DİLEKÇESİ").bold = True
        
        # Boşluk
        doc.add_paragraph()
        
        # Karar bilgileri
        karar = doc.add_paragraph()
        karar.alignment = WD_ALIGN_PARAGRAPH.LEFT
        karar.add_run("TEMYİZ EDİLEN KARAR: ").bold = True
        karar.add_run(template_data.get("karar_bilgisi", ""))
        
        # Taraf bilgileri
        temyiz_eden = doc.add_paragraph()
        temyiz_eden.alignment = WD_ALIGN_PARAGRAPH.LEFT
        temyiz_eden.add_run("TEMYİZ EDEN: ").bold = True
        temyiz_eden.add_run(template_data.get("temyiz_eden", ""))
        
        # Karşı taraf bilgileri
        karsi_taraf = doc.add_paragraph()
        karsi_taraf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        karsi_taraf.add_run("KARŞI TARAF: ").bold = True
        karsi_taraf.add_run(template_data.get("karsi_taraf", ""))
        
        # Temyiz sebepleri
        doc.add_paragraph()
        temyiz_sebepleri_baslik = doc.add_paragraph()
        temyiz_sebepleri_baslik.alignment = WD_ALIGN_PARAGRAPH.CENTER
        temyiz_sebepleri_baslik.add_run("TEMYİZ SEBEPLERİ").bold = True
        
        temyiz_sebepleri = doc.add_paragraph()
        temyiz_sebepleri.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        temyiz_sebepleri.add_run(template_data.get("temyiz_sebepleri", ""))
        
        # Sonuç ve Talep
        doc.add_paragraph()
        sonuc_talep = doc.add_paragraph()
        sonuc_talep.alignment = WD_ALIGN_PARAGRAPH.LEFT
        sonuc_talep.add_run("SONUÇ VE TALEP: ").bold = True
        sonuc_talep.add_run(template_data.get("sonuc_talep", ""))
        
        # Tarih ve imza
        doc.add_paragraph()
        signature = doc.add_paragraph()
        signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date = datetime.now().strftime("%d/%m/%Y")
        signature.add_run(f"Tarih: {date}\n\n")
        signature.add_run(f"Ad Soyad: {template_data.get('ad_soyad', '')}\n\n")
        signature.add_run("İmza: _________________")
        
        # Ekler
        if "ekler" in template_data and template_data["ekler"]:
            doc.add_paragraph()
            ekler = doc.add_paragraph()
            ekler.add_run("EKLER:\n").bold = True
            for i, ek in enumerate(template_data["ekler"], 1):
                ekler.add_run(f"{i}- {ek}\n")
        
        # Belgeyi byte olarak kaydet
        import io
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()

class TemplateManager:
    def get_available_templates(self) -> List[TemplateInfo]:
        """Kullanılabilir tüm şablonları döndürür"""
        templates = [
            TemplateInfo(
                id="dilekce",
                name="Dilekçe",
                description="Genel amaçlı dilekçe şablonu",
                category="Genel",
                required_fields=[
                    {"name": "kurum", "type": "text", "description": "Dilekçenin sunulacağı kurum"},
                    {"name": "konu", "type": "text", "description": "Dilekçe konusu"},
                    {"name": "icerik", "type": "text", "description": "Dilekçe içeriği"},
                    {"name": "ad_soyad", "type": "text", "description": "Ad Soyad"},
                    {"name": "ekler", "type": "list", "description": "Ekler listesi", "optional": True}
                ]
            ),
            TemplateInfo(
                id="ihtarname",
                name="İhtarname",
                description="Resmi uyarı belgesi şablonu",
                category="Uyarı",
                required_fields=[
                    {"name": "gonderen", "type": "text", "description": "İhtarı gönderen kişi/kurum bilgileri"},
                    {"name": "alici", "type": "text", "description": "İhtarın alıcısı kişi/kurum bilgileri"},
                    {"name": "konu", "type": "text", "description": "İhtar konusu"},
                    {"name": "icerik", "type": "text", "description": "İhtar detayları"},
                    {"name": "sonuc_talep", "type": "text", "description": "Sonuç ve talep"},
                    {"name": "ad_soyad", "type": "text", "description": "Ad Soyad"}
                ]
            ),
            TemplateInfo(
                id="vekaletname",
                name="Vekaletname",
                description="Vekil tayini için kullanılan belge şablonu",
                category="Yetkilendirme",
                required_fields=[
                    {"name": "vekil_eden", "type": "text", "description": "Vekil eden kişi/kurum bilgileri"},
                    {"name": "vekil", "type": "text", "description": "Vekil kişi bilgileri"},
                    {"name": "icerik", "type": "text", "description": "Vekaletname içeriği"},
                    {"name": "yetkiler", "type": "list", "description": "Verilen yetkiler listesi"},
                    {"name": "ad_soyad", "type": "text", "description": "Ad Soyad"}
                ]
            ),
            TemplateInfo(
                id="dava_dilekce",
                name="Dava Dilekçesi",
                description="Hukuk davası açmak için kullanılan dilekçe şablonu",
                category="Dava",
                required_fields=[
                    {"name": "mahkeme", "type": "text", "description": "Mahkeme adı"},
                    {"name": "dava_turu", "type": "text", "description": "Dava türü"},
                    {"name": "davaci", "type": "text", "description": "Davacı bilgileri"},
                    {"name": "davali", "type": "text", "description": "Davalı bilgileri"},
                    {"name": "konu", "type": "text", "description": "Dava konusu"},
                    {"name": "deger", "type": "text", "description": "Dava değeri"},
                    {"name": "aciklamalar", "type": "text", "description": "Dava açıklamaları"},
                    {"name": "hukuki_sebepler", "type": "text", "description": "Hukuki sebepler"},
                    {"name": "deliller", "type": "list", "description": "Deliller listesi"},
                    {"name": "sonuc_talep", "type": "text", "description": "Sonuç ve talep"},
                    {"name": "davaci_ad_soyad", "type": "text", "description": "Davacı Ad Soyad"},
                    {"name": "ekler", "type": "list", "description": "Ekler listesi", "optional": True}
                ]
            ),
            TemplateInfo(
                id="temyiz_dilekce",
                name="Temyiz Dilekçesi",
                description="Karar temyizi için kullanılan dilekçe şablonu",
                category="Temyiz",
                required_fields=[
                    {"name": "mahkeme", "type": "text", "description": "Yüksek Mahkeme adı"},
                    {"name": "karar_bilgisi", "type": "text", "description": "Temyiz edilen karar bilgileri"},
                    {"name": "temyiz_eden", "type": "text", "description": "Temyiz eden taraf bilgileri"},
                    {"name": "karsi_taraf", "type": "text", "description": "Karşı taraf bilgileri"},
                    {"name": "temyiz_sebepleri", "type": "text", "description": "Temyiz sebepleri"},
                    {"name": "sonuc_talep", "type": "text", "description": "Sonuç ve talep"},
                    {"name": "ad_soyad", "type": "text", "description": "Ad Soyad"},
                    {"name": "ekler", "type": "list", "description": "Ekler listesi", "optional": True}
                ]
            )
        ]
        return templates
        
    def validate_template_data(self, template_name: str, template_data: Dict[str, Any]) -> list:
        """Şablon verilerini doğrula ve hataları liste olarak döndür"""
        errors = []
        
        # Şablonlar için gerekli alanlar
        required_fields = {
            "dilekce": ["kurum", "konu", "icerik", "ad_soyad"],
            "ihtarname": ["gonderen", "alici", "konu", "icerik", "sonuc_talep", "ad_soyad"],
            "vekaletname": ["vekil_eden", "vekil", "icerik", "yetkiler", "ad_soyad"],
            "dava_dilekce": ["mahkeme", "dava_turu", "davaci", "davali", "konu", "deger", 
                             "aciklamalar", "hukuki_sebepler", "deliller", "sonuc_talep", "davaci_ad_soyad"],
            "temyiz_dilekce": ["mahkeme", "karar_bilgisi", "temyiz_eden", "karsi_taraf", 
                               "temyiz_sebepleri", "sonuc_talep", "ad_soyad"]
        }
        
        # Şablon mevcut mu kontrol et
        if template_name not in required_fields:
            errors.append(f"Geçersiz şablon adı: {template_name}")
            return errors
            
        # Gerekli alanlar mevcut mu kontrol et
        for field in required_fields[template_name]:
            if field not in template_data or not template_data[field]:
                errors.append(f"Gerekli alan eksik: {field}")
                
        return errors

# Bileşenleri başlat
doc_generator = DocumentGenerator()
template_manager = TemplateManager()

@app.get("/templates", response_model=TemplateListResponse)
async def list_templates():
    """Kullanılabilir tüm şablonları listele"""
    templates = template_manager.get_available_templates()
    return TemplateListResponse(templates=templates)

@app.get("/templates/{template_id}")
async def get_template(template_id: str):
    """Belirli bir şablonun ayrıntılarını getir"""
    templates = template_manager.get_available_templates()
    template = next((t for t in templates if t.id == template_id), None)
    if not template:
        raise HTTPException(status_code=404, detail="Şablon bulunamadı")
    return template

@app.get("/documents/{doc_id}/download")
async def download_document(doc_id: str) -> FileResponse:
    """Oluşturulan belgeyi indirme"""
    try:
        # Belge ID formatını doğrula
        try:
            uuid.UUID(doc_id)
        except ValueError:
            raise HTTPException(status_code=400, detail="Geçersiz belge ID formatı")

        # Belgeyi getir
        doc_path = f"storage/documents/{doc_id}.docx"
        if not os.path.exists(doc_path):
            raise HTTPException(status_code=404, detail="Belge bulunamadı")
            
        return FileResponse(
            doc_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"belge_{doc_id}.docx"
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/generate", response_model=GenerateResponse)
async def generate_document(
    request: GenerateDocumentRequest,
    background_tasks: BackgroundTasks
) -> GenerateResponse:
    """Yeni bir hukuki belge oluştur"""
    try:
        # Şablon verilerini doğrula
        errors = template_manager.validate_template_data(
            request.template_name,
            request.template_data
        )
        if errors:
            raise HTTPException(
                status_code=400,
                detail={"message": "Geçersiz şablon verileri", "errors": errors}
            )
        
        # Belgeyi oluştur
        doc_content = doc_generator.generate_document(
            request.template_name, 
            request.template_data
        )
        
        # Belge ID'si oluştur
        doc_id = str(uuid.uuid4())
        
        # Belgeyi kaydet
        background_tasks.add_task(store_generated_document, doc_id, doc_content)
        
        # Metadata sözlüğü oluştur
        metadata = {
            "template": request.template_name,
            "generated_at": datetime.now().isoformat(),
            "size": len(doc_content)
        }
        if request.metadata:
            metadata.update(request.metadata)
        
        return GenerateResponse(
            document_id=doc_id,
            download_url=f"/documents/{doc_id}/download",
            metadata=metadata
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/legal-advice/categories")
async def get_legal_advice_categories():
    """Hukuki danışmanlık kategorilerini listele"""
    categories = [
        {"id": "iş_hukuku", "name": "İş Hukuku"},
        {"id": "aile_hukuku", "name": "Aile Hukuku"},
        {"id": "borçlar_hukuku", "name": "Borçlar Hukuku"},
        {"id": "ceza_hukuku", "name": "Ceza Hukuku"},
        {"id": "tüketici_hukuku", "name": "Tüketici Hukuku"},
    ]
    return {"categories": categories}

@app.get("/legal-advice/{category}")
async def get_legal_advice(
    category: str, 
    query: str = Query(None, description="Hukuki soru içeriği")
):
    """Belirli bir kategoride hukuki tavsiye sağla"""
    # Bu kısım daha sonra AI tabanlı gerçek bir tavsiye sistemi ile değiştirilebilir
    # Şimdilik basit statik yanıtlar döndürelim
    
    advice_responses = {
        "iş_hukuku": {
            "default": "İş hukuku, işçi ve işveren arasındaki ilişkileri düzenleyen hukuk dalıdır. İş sözleşmesinin türleri, işten çıkarma koşulları, kıdem tazminatı, ihbar tazminatı gibi konuları kapsar."
        },
        "aile_hukuku": {
            "default": "Aile hukuku; evlilik, boşanma, velayet, nafaka, mal paylaşımı gibi aile ilişkilerini düzenleyen hukuk dalıdır."
        },
        "borçlar_hukuku": {
            "default": "Borçlar hukuku, kişiler arasındaki borç ilişkilerini düzenleyen özel hukuk dalıdır. Sözleşmeler, haksız fiiller, sebepsiz zenginleşme gibi konuları kapsar."
        },
        "ceza_hukuku": {
            "default": "Ceza hukuku, suç teşkil eden fiilleri ve bunlara uygulanacak yaptırımları belirleyen hukuk dalıdır."
        },
        "tüketici_hukuku": {
            "default": "Tüketici hukuku, tüketiciler ile satıcılar arasındaki ilişkileri düzenleyen ve tüketiciyi koruyan hukuk dalıdır."
        }
    }
    
    if category not in advice_responses:
        raise HTTPException(status_code=404, detail="Kategori bulunamadı")
    
    # Basit bir yanıt döndür
    response = {
        "category": category,
        "query": query if query else "Genel bilgi",
        "advice": advice_responses[category]["default"]
    }
    
    return response

async def store_generated_document(doc_id: str, content: bytes):
    """Oluşturulan belgeyi kaydet"""
    try:
        os.makedirs("storage/documents", exist_ok=True)
        file_path = f"storage/documents/{doc_id}.docx"
        with open(file_path, "wb") as f:
            f.write(content)
    except Exception as e:
        # Hata kaydı tutulabilir
        raise Exception(f"Belge {doc_id} kaydedilemedi: {str(e)}")

@app.get("/", response_class=FileResponse)
async def read_index():
    """Ana sayfayı göster"""
    return FileResponse("app/static/index.html")

@app.post("/analyze-case", response_model=CaseAnalysisResponse)
async def analyze_case(request: CaseAnalysisRequest) -> CaseAnalysisResponse:
    """
    Hukuki bir vakayı analiz eder ve ilgili kanunları ve Yargıtay kararlarını bulur.
    """
    try:
        # AI analizi yap
        legal_analysis = await ai_service.analyze_case(
            request.case_description, 
            request.case_category
        )
        
        # İlgili kanunlar ve kararları bul
        relevant_data = await legal_db_service.get_relevant_laws_and_decisions(
            request.case_description, 
            request.case_category
        )
        
        # Analiz metninden kanun ve karar referanslarını çıkar ve veritabanından detaylarını al
        enhanced_results = await extract_and_enhance_references(legal_analysis, legal_db_service)
        
        # Sonuçları birleştir
        all_laws = relevant_data["laws"] + enhanced_results["laws"]
        all_decisions = relevant_data["court_decisions"] + enhanced_results["court_decisions"]
        
        # Tekrarlanan kayıtları kaldır
        unique_laws = []
        law_ids = set()
        for law in all_laws:
            if law.get("id") not in law_ids:
                unique_laws.append(law)
                law_ids.add(law.get("id"))
        
        unique_decisions = []
        decision_ids = set()
        for decision in all_decisions:
            if decision.get("id") not in decision_ids:
                unique_decisions.append(decision)
                decision_ids.add(decision.get("id"))
        
        return CaseAnalysisResponse(
            analysis=legal_analysis,
            relevant_laws=unique_laws,
            relevant_decisions=unique_decisions
        )
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Analiz sırasında bir hata oluştu: {str(e)}"
        )

async def extract_and_enhance_references(legal_analysis: Dict[str, Any], db_service: LegalDatabaseService) -> Dict[str, List]:
    """
    Yapay zeka analizi metninden kanun ve Yargıtay kararı referanslarını çıkarır
    ve veritabanında arayarak ayrıntılı bilgilerini alır.
    """
    result = {
        "laws": [],
        "court_decisions": []
    }
    
    if not legal_analysis or "analysis" not in legal_analysis:
        return result
    
    analysis_text = legal_analysis["analysis"]
    
    # Kanun referanslarını çıkar (TMK 134, TBK 49 gibi)
    law_refs_pattern = r'\b(TMK|TBK|TCK|İK|HMK|TTK|TKHK|İYUK|HUMK|İİK)\s+(\d+)'
    
    law_code_map = {
        "TMK": "4721",  # Türk Medeni Kanunu
        "TBK": "6098",  # Türk Borçlar Kanunu
        "TCK": "5237",  # Türk Ceza Kanunu
        "İK": "4857",   # İş Kanunu
        "HMK": "6100",  # Hukuk Muhakemeleri Kanunu
        "TTK": "6102",  # Türk Ticaret Kanunu
        "TKHK": "6502", # Tüketicinin Korunması Hakkında Kanun
        "İYUK": "2577", # İdari Yargılama Usulü Kanunu
        "İİK": "2004",  # İcra ve İflas Kanunu
    }
    
    law_refs = re.findall(law_refs_pattern, analysis_text)
    processed_laws = set()
    
    for law_abbr, article_no in law_refs:
        if law_abbr in law_code_map:
            law_no = law_code_map[law_abbr]
            
            # Aynı kanunu tekrar işlememek için kontrol
            if law_no in processed_laws:
                continue
                
            # Kanunu veritabanında ara
            conn = sqlite3.connect(db_service.db_path)
            cursor = conn.cursor()
            
            cursor.execute('''
            SELECT id, law_no, name, category, content, publication_date, last_updated
            FROM laws
            WHERE law_no = ?
            LIMIT 1
            ''', (law_no,))
            
            row = cursor.fetchone()
            
            if row:
                result["laws"].append({
                    "id": row[0],
                    "law_no": row[1],
                    "name": row[2],
                    "category": row[3],
                    "content": row[4],
                    "publication_date": row[5],
                    "last_updated": row[6]
                })
                processed_laws.add(law_no)
                
                # İlgili maddeyi de ara
                cursor.execute('''
                SELECT id, law_id, article_no, content
                FROM law_articles
                WHERE law_id = ? AND article_no = ?
                LIMIT 1
                ''', (row[0], article_no))
                
                article_row = cursor.fetchone()
                if article_row:
                    # Madde içeriğini eğer bulunabilirse konsolda göster, log için
                    print(f"Madde bulundu: {row[2]} Madde {article_no}, İçerik: {article_row[3][:100]}...")
            
            conn.close()
    
    # Yargıtay kararlarını çıkar (Yargıtay 9. Hukuk Dairesi 2019/1234 gibi)
    decision_pattern = r'Yargıtay\s+(\d+)\.\s+(Hukuk|Ceza)\s+Dairesi\s*(?:[^0-9]*?)\s*(\d{4})[\/\-](\d+)'
    
    decision_refs = re.findall(decision_pattern, analysis_text)
    processed_decisions = set()
    
    for chamber_no, chamber_type, year, number in decision_refs:
        decision_id = f"{year}/{number}"
        chamber = f"Yargıtay {chamber_no}. {chamber_type} Dairesi"
        
        # Aynı kararı tekrar işlememek için kontrol
        if decision_id in processed_decisions:
            continue
        
        # Kararı veritabanında ara
        conn = sqlite3.connect(db_service.db_path)
        cursor = conn.cursor()
        
        cursor.execute('''
        SELECT id, decision_no, decision_date, chamber, subject, content, keywords
        FROM court_decisions
        WHERE decision_no = ? OR (chamber LIKE ? AND decision_no LIKE ?)
        LIMIT 1
        ''', (decision_id, f"%{chamber}%", f"%{year}%{number}%"))
        
        row = cursor.fetchone()
        
        if row:
            result["court_decisions"].append({
                "id": row[0],
                "decision_no": row[1],
                "decision_date": row[2],
                "chamber": row[3],
                "subject": row[4],
                "content": row[5],
                "keywords": row[6]
            })
            processed_decisions.add(decision_id)
        
        conn.close()
    
    return result

@app.post("/generate-ai-enhanced", response_model=GenerateResponse)
async def generate_ai_enhanced_document(
    request: GenerateAIEnhancedDocumentRequest,
    background_tasks: BackgroundTasks
) -> GenerateResponse:
    """
    Yapay zeka destekli gelişmiş bir hukuki belge oluşturur.
    """
    try:
        print(f"[DEBUG] Received request for template: {request.template_name}, category: {request.case_category}")
        
        # AI tarafından doldurulacak alanları önceden geçici değerlerle doldur
        # Bu, doğrulama aşamasından geçmesini sağlayacak
        template_data = request.template_data.copy()
        
        # Şablon türüne göre gerekli alanları belirle ve geçici değerler ata
        if request.template_name == "dilekce":
            if "icerik" not in template_data:
                template_data["icerik"] = "[AI tarafından doldurulacak]"
        elif request.template_name == "ihtarname":
            for field in ["icerik", "sonuc_talep"]:
                if field not in template_data:
                    template_data[field] = "[AI tarafından doldurulacak]"
        elif request.template_name == "dava_dilekce":
            for field in ["aciklamalar", "hukuki_sebepler", "sonuc_talep"]:
                if field not in template_data:
                    template_data[field] = "[AI tarafından doldurulacak]"
        elif request.template_name == "temyiz_dilekce":
            for field in ["temyiz_sebepleri", "sonuc_talep"]:
                if field not in template_data:
                    template_data[field] = "[AI tarafından doldurulacak]"
        elif request.template_name == "vekaletname":
            if "yetki_kapsami" not in template_data:
                template_data["yetki_kapsami"] = "[AI tarafından doldurulacak]"
                
        # Güncellenmiş şablon verilerini doğrula
        errors = template_manager.validate_template_data(
            request.template_name,
            template_data
        )
        if errors:
            print(f"[ERROR] Template validation failed: {errors}")
            raise HTTPException(
                status_code=400,
                detail={"message": "Geçersiz şablon verileri", "errors": errors}
            )
        
        # AI analizi yap
        print(f"[DEBUG] Starting AI analysis for case description (length: {len(request.case_description)})")
        legal_analysis = await ai_service.analyze_case(
            request.case_description, 
            request.case_category
        )
        
        if "error" in legal_analysis:
            print(f"[ERROR] AI analysis failed: {legal_analysis['error']}")
            legal_analysis = {
                "analysis": "Analiz yapılamadı, ancak belge oluşturulabilir.",
                "category": request.case_category,
                "relevant_laws": [],
                "relevant_decisions": []
            }
        
        # AI destekli belge içeriği oluştur
        print(f"[DEBUG] Generating document content with AI for template: {request.template_name}")
        enhanced_document_data = await ai_service.generate_legal_document(
            request.template_name,
            request.case_description,
            legal_analysis,
            template_data  # Geçici değerlerle doldurulmuş şablon verilerini kullan
        )
        
        # Hata kontrolü
        if "error" in enhanced_document_data:
            print(f"[ERROR] Document generation failed: {enhanced_document_data['error']}")
            # Hata olsa bile, basit bir içerik ile devam et
            document_data = template_data.copy()
            document_data["icerik"] = f"Bu belge için içerik oluşturulamadı. Lütfen tekrar deneyiniz.\n\nOlay: {request.case_description[:200]}..."
            enhanced_document_data["enhanced_template_data"] = document_data
        
        # Belgeyi oluştur
        print(f"[DEBUG] Generating document with template data")
        try:
            doc_content = doc_generator.generate_document(
                request.template_name, 
                enhanced_document_data["enhanced_template_data"]
            )
            print(f"[DEBUG] Document generation successful, content size: {len(doc_content)}")
        except Exception as doc_ex:
            print(f"[ERROR] Error in document generator: {str(doc_ex)}")
            raise HTTPException(
                status_code=500,
                detail=f"Belge oluşturma hatası: {str(doc_ex)}"
            )
        
        # Belge ID'si oluştur
        doc_id = str(uuid.uuid4())
        
        # Belgeyi kaydet
        try:
            background_tasks.add_task(store_generated_document, doc_id, doc_content)
            print(f"[DEBUG] Document scheduled for storage with ID: {doc_id}")
        except Exception as storage_ex:
            print(f"[ERROR] Error storing document: {str(storage_ex)}")
            # Storage hatası olsa da devam et
        
        # Metadata
        metadata = {
            "template": request.template_name,
            "generated_at": datetime.now().isoformat(),
            "size": len(doc_content),
            "ai_enhanced": True,
            "legal_analysis": {
                "category": request.case_category,
                "timestamp": legal_analysis.get("timestamp", datetime.now().isoformat())
            }
        }
        
        if request.metadata:
            metadata.update(request.metadata)
        
        print(f"[DEBUG] Returning successful response for document ID: {doc_id}")
        return GenerateResponse(
            document_id=doc_id,
            download_url=f"/documents/{doc_id}/download",
            metadata=metadata
        )
    except ValueError as e:
        print(f"[ERROR] Value error in generate_ai_enhanced: {str(e)}")
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        print(f"[ERROR] Unexpected error in generate_ai_enhanced: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/legal-database/laws")
async def search_laws(
    query: str = Query(..., description="Arama sorgusu"),
    category: str = Query(None, description="Kanun kategorisi")
):
    """
    Kanunlar veritabanında arama yapar.
    """
    try:
        results = await legal_db_service.search_laws(query, category)
        return {"results": results, "count": len(results)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/legal-database/court-decisions")
async def search_court_decisions(
    query: str = Query(..., description="Arama sorgusu"),
    chamber: str = Query(None, description="Yargıtay dairesi")
):
    """
    Yargıtay kararları veritabanında arama yapar.
    """
    try:
        results = await legal_db_service.search_court_decisions(query, chamber)
        return {"results": results, "count": len(results)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/initialize-database")
async def initialize_database(background_tasks: BackgroundTasks):
    """
    Veritabanını örnek verilerle doldurur (sadece test için).
    """
    try:
        background_tasks.add_task(legal_db_service.populate_sample_data)
        return {"message": "Veritabanı doldurma görevi başlatıldı"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/data-collector/scrape-laws")
async def scrape_laws(
    law_type: str = Query("kanun", description="Mevzuat tipi (kanun, khk, tüzük, yönetmelik vb.)"),
    limit: int = Query(50, description="Toplanacak maksimum kanun sayısı"),
    background_tasks: BackgroundTasks = None
):
    """
    Mevzuat.gov.tr sitesinden kanunları çekerek veritabanına kaydeder.
    Bu endpoint, kanunları arka planda toplar.
    """
    # Bu işlemi arka planda yap
    if background_tasks:
        background_tasks.add_task(
            data_collector_service.collect_laws_from_mevzuat,
            law_type=law_type,
            limit=limit
        )
        return {"status": "success", "message": f"{limit} adet {law_type} toplanıyor. Bu işlem arka planda devam edecek."}
    else:
        # Eğer background_tasks yoksa senkron olarak çalıştır
        result = await data_collector_service.collect_laws_from_mevzuat(law_type, limit)
        return {"status": "success", "message": f"{result} adet {law_type} toplandı."}

@app.post("/api/data-collector/scrape-court-decisions")
async def scrape_court_decisions(
    chamber: str = Query(None, description="Yargıtay dairesi (örn: '9. Hukuk Dairesi')"),
    start_date: str = Query(None, description="Başlangıç tarihi (YYYY-MM-DD formatında)"),
    end_date: str = Query(None, description="Bitiş tarihi (YYYY-MM-DD formatında)"),
    limit: int = Query(50, description="Toplanacak maksimum karar sayısı"),
    background_tasks: BackgroundTasks = None
):
    """
    Yargıtay kararlarını çekerek veritabanına kaydeder.
    Bu endpoint, kararları arka planda toplar.
    """
    # Bu işlemi arka planda yap
    if background_tasks:
        background_tasks.add_task(
            data_collector_service.collect_court_decisions,
            chamber=chamber,
            start_date=start_date, 
            end_date=end_date,
            limit=limit
        )
        return {"status": "success", "message": f"Yargıtay kararları toplanıyor. Bu işlem arka planda devam edecek."}
    else:
        # Eğer background_tasks yoksa senkron olarak çalıştır
        result = await data_collector_service.collect_court_decisions(chamber, start_date, end_date, limit)
        return {"status": "success", "message": f"{result} adet Yargıtay kararı toplandı."}

@app.post("/api/data-collector/backup-database")
async def backup_database(
    backup_path: str = Query(None, description="Yedek dosyasının yolu")
):
    """
    Veritabanını yedekler.
    """
    result = await data_collector_service.backup_database(backup_path)
    if result:
        return {"status": "success", "message": "Veritabanı başarıyla yedeklendi."}
    else:
        raise HTTPException(
            status_code=500,
            detail="Veritabanı yedeklenirken bir hata oluştu."
        )

@app.get("/legal-database/laws/{law_id}")
async def get_law_details(law_id: int):
    """
    Belirli bir kanunun detaylarını döndürür.
    """
    try:
        conn = sqlite3.connect(legal_db_service.db_path)
        cursor = conn.cursor()
        
        # Kanun bilgilerini al
        cursor.execute('''
        SELECT id, law_no, name, category, content, publication_date, last_updated
        FROM laws
        WHERE id = ?
        ''', (law_id,))
        
        law = cursor.fetchone()
        if not law:
            raise HTTPException(status_code=404, detail="Kanun bulunamadı")
            
        # Kanun maddelerini al
        cursor.execute('''
        SELECT id, article_no, content
        FROM law_articles
        WHERE law_id = ?
        ORDER BY article_no
        ''', (law_id,))
        
        articles = []
        for article in cursor.fetchall():
            articles.append({
                "id": article[0],
                "article_no": article[1],
                "content": article[2]
            })
        
        conn.close()
        
        return {
            "id": law[0],
            "law_no": law[1],
            "name": law[2],
            "category": law[3],
            "content": law[4],
            "publication_date": law[5],
            "last_updated": law[6],
            "articles": articles
        }
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Kanun detayları alınırken bir hata oluştu: {str(e)}"
        )

@app.get("/legal-database/court-decisions/{decision_id}")
async def get_court_decision_details(decision_id: int):
    """
    Belirli bir Yargıtay kararının detaylarını döndürür.
    """
    try:
        conn = sqlite3.connect(legal_db_service.db_path)
        cursor = conn.cursor()
        
        cursor.execute('''
        SELECT id, decision_no, decision_date, chamber, subject, content, keywords
        FROM court_decisions
        WHERE id = ?
        ''', (decision_id,))
        
        decision = cursor.fetchone()
        if not decision:
            raise HTTPException(status_code=404, detail="Yargıtay kararı bulunamadı")
        
        conn.close()
        
        return {
            "id": decision[0],
            "decision_no": decision[1],
            "decision_date": decision[2],
            "chamber": decision[3],
            "subject": decision[4],
            "content": decision[5],
            "keywords": decision[6]
        }
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Yargıtay kararı detayları alınırken bir hata oluştu: {str(e)}"
        )

@app.post("/reset-database")
async def reset_database():
    """
    Veritabanını sıfırlar ve tabloları yeniden oluşturur.
    Dikkat: Bu işlem tüm verileri siler.
    """
    try:
        from app.services.legal_database_service import LegalDatabaseService
        
        db_service = LegalDatabaseService()
        success = await db_service.reset_database()
        
        if success:
            return {"status": "success", "message": "Veritabanı başarıyla sıfırlandı"}
        else:
            return {"status": "error", "message": "Veritabanı sıfırlanamadı"}
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"status": "error", "message": f"Hata: {str(e)}"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
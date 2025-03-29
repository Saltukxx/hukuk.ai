"""
Hukuk.AI - Legal Document Generator API
A simplified FastAPI application for generating legal documents
"""

from fastapi import FastAPI, HTTPException, Request
from fastapi.responses import FileResponse, JSONResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Dict, Any, Optional, List
import os
import json
import uuid
from datetime import datetime
import logging
from dotenv import load_dotenv
from app.services.document_generator import DocumentGenerator
from app.services.ai_service import AILegalAnalyzer
from app.models import DocumentRequest, DocumentResponse, AIDocumentRequest, LegalAnalysis
from app.core.config import get_gemini_api_key, API_USE_MOCK_DATA

# Load environment variables directly here as well to ensure they're available
load_dotenv()

# Force API_USE_MOCK_DATA to be False to ensure real AI analysis is used
FORCE_REAL_AI = True  # This will override the config setting
MOCK_DATA_ENABLED = False  # Additional safety to disable mock data

# Configure logging with more detailed format
logging.basicConfig(
    level=logging.DEBUG if os.getenv("DEBUG", "False").lower() in ("true", "1", "t") else logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s - %(filename)s:%(lineno)d",
    handlers=[
        logging.FileHandler("app/logs/app.log"),
        logging.StreamHandler()
    ]
)

logger = logging.getLogger(__name__)

# Verify API key
api_key = get_gemini_api_key()
if api_key:
    logger.info(f"API key found in environment, length: {len(api_key)}")
    masked_key = api_key[:4] + "..." + api_key[-4:] if len(api_key) > 8 else "[masked]"
    logger.info(f"API key format: {masked_key}")
else:
    logger.error("⚠️ NO API KEY FOUND! Application will not be able to use real AI analysis!")

# Log the configuration for AI analysis
logger.info(f"FORCE_REAL_AI set to: {FORCE_REAL_AI}")
logger.info(f"MOCK_DATA_ENABLED set to: {MOCK_DATA_ENABLED}")
logger.info(f"API_USE_MOCK_DATA from config: {API_USE_MOCK_DATA}")

# Initialize FastAPI app
app = FastAPI(
    title="Hukuk.AI - Türkçe Hukuki Belge Üretici",
    description="Dilekçe ve hukuki belge hazırlama sistemi",
    version="1.0.0"
)

# CORS settings
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allow all origins for development
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize the document generator service
document_generator = DocumentGenerator()

# Initialize the AI legal analyzer service with the API key from config
ai_legal_analyzer = AILegalAnalyzer(api_key=api_key)

# Create the output directory if it doesn't exist
os.makedirs("app/output", exist_ok=True)
os.makedirs("app/logs", exist_ok=True)

# Define request/response models
class AIDocumentRequest(BaseModel):
    template_name: str
    case_description: str
    case_category: str
    template_data: Dict[str, Any]
    metadata: Optional[Dict[str, Any]] = None

class DocumentMetadata(BaseModel):
    category: str
    keywords: List[str] = []
    language: str = "tr"

class DocumentResponse(BaseModel):
    document_id: str
    title: str
    document_type: str
    content: str
    metadata: Dict[str, Any]
    analysis: Optional[Dict[str, Any]] = None

# Mount static files after all other routes to prevent conflicts
# This is important because the order matters in FastAPI
@app.get("/")
async def root():
    """
    Serve the main index.html file
    """
    # Explicitly return the index.html file
    try:
        with open("app/static/index.html", "r", encoding="utf-8") as f:
            content = f.read()
        return HTMLResponse(content=content)
    except Exception as e:
        logger.error(f"Error serving index.html: {str(e)}")
        return {"error": "Could not serve the index page"}

# Mount static files for other assets
app.mount("/static", StaticFiles(directory="app/static"), name="static")

@app.post("/api/documents/ai-generate", response_model=DocumentResponse)
async def generate_document(request: AIDocumentRequest):
    """
    Generate a document based on AI analysis of the case description
    
    Args:
        request: The document generation request
        
    Returns:
        DocumentResponse: The generated document details and analysis
    """
    try:
        # Log the request
        logger.info(f"Received document generation request: Template={request.template_name}, Category={request.case_category}")
        description_preview = request.case_description[:100] + "..." if len(request.case_description) > 100 else request.case_description
        logger.info(f"Case description preview: {description_preview}")
        
        # Create a unique document ID
        document_id = str(uuid.uuid4())
        logger.info(f"Created document ID: {document_id}")
        
        # Initialize the document generator service
        generator = DocumentGenerator()
        logger.info("DocumentGenerator initialized")
        
        # Get the method that will generate the document
        generate_method = getattr(generator, "generate_document", None)
        
        if generate_method:
            # Log the parameters for debugging
            import inspect
            sig = inspect.signature(generate_method)
            logger.info(f"Found generate_document method with parameters: {list(sig.parameters.keys())}")
            
            # Check if the method expects document_id
            if "document_id" in sig.parameters:
                logger.info("Calling generate_document with document_id")
                result = await generate_method(
                    template_name=request.template_name,
                    template_data=request.template_data,
                    document_id=document_id,
                    output_format="docx"
                )
            else:
                logger.info("Calling generate_document without document_id")
                result = await generate_method(
                    template_name=request.template_name,
                    template_data=request.template_data,
                    output_format="docx"
                )
                
            # Extract the file path from the result
            file_path = result.get("document_path") if isinstance(result, dict) else result
        else:
            logger.error("Generate document method not found")
            raise HTTPException(status_code=500, detail="Document generation method not found")
            
        # If the file was generated successfully, log it
        if file_path and os.path.exists(file_path):
            logger.info(f"Successfully generated document at: {file_path}")
            
            # Convert DOCX to HTML for direct display
            document_html = await convert_docx_to_html(file_path)
            logger.info(f"Converted document to HTML: {len(document_html)} characters")
        else:
            logger.error(f"Failed to generate document, file path not found: {file_path}")
            raise HTTPException(status_code=500, detail="Document generation failed - output file not found")
        
        # Use the AI Legal Analyzer to analyze the case and get relevant laws/decisions
        logger.info(f"Requesting AI analysis for case in category: {request.case_category}")
        
        # First, try to get AI analysis with retry mechanism
        max_retries = 3  # Increased from 2 to 3 for more attempts
        using_mock_data = False
        analysis_data = None
        
        for attempt in range(max_retries):
            try:
                logger.info(f"AI analysis attempt {attempt+1}/{max_retries} - FORCING REAL ANALYSIS")
                # Get AI analysis for the case - force real AI analysis 
                analysis_data = await ai_legal_analyzer.analyze_case(
                    case_description=request.case_description,
                    case_category=request.case_category,
                    force_real_analysis=True  # Always force real analysis
                )
                
                # Verify we have actual content (not empty sections)
                has_laws = len(analysis_data.get("relevant_laws", [])) > 0
                has_decisions = len(analysis_data.get("relevant_decisions", [])) > 0
                
                if has_laws and has_decisions:
                    logger.info(f"✅ AI analysis completed successfully with content - Laws: {len(analysis_data.get('relevant_laws', []))}, Decisions: {len(analysis_data.get('relevant_decisions', []))}")
                    break
                else:
                    logger.warning(f"⚠️ AI returned empty or incomplete analysis, will retry. Has laws: {has_laws}, Has decisions: {has_decisions}")
                    
            except Exception as analysis_error:
                logger.error(f"Error during AI analysis attempt {attempt+1}: {str(analysis_error)}")
                logger.exception(analysis_error)
        
        # Only fall back to mock data if real analysis completely fails and if mock data is allowed
        if (not analysis_data or (not len(analysis_data.get("relevant_laws", [])) and not len(analysis_data.get("relevant_decisions", [])))) and MOCK_DATA_ENABLED:
            logger.warning("Falling back to mock data for analysis after failed attempts - THIS SHOULD NOT HAPPEN IN PRODUCTION")
            analysis_data = ai_legal_analyzer._get_mock_analysis(request.case_category)
            using_mock_data = True
        elif not analysis_data:
            logger.error("AI analysis failed and mock data is not allowed or disabled")
            analysis_data = {
                "summary": "AI analizi başarısız oldu",
                "relevant_laws": [
                    {
                        "title": "Analiz Hatası",
                        "description": "Yapay zeka analizi başarısız oldu. Lütfen daha ayrıntılı bir olay özeti yazarak tekrar deneyin."
                    }
                ],
                "relevant_decisions": [
                    {
                        "case_number": "Hata",
                        "date": "Belirsiz",
                        "summary": "Yapay zeka analizi başarısız oldu. Lütfen olay özetini genişleterek tekrar deneyin."
                    }
                ],
                "recommendations": "Lütfen olay özetinizi daha ayrıntılı açıklayarak tekrar deneyin. En az 200 karakter içeren detaylı bir açıklama, yapay zeka analizinin daha doğru olmasını sağlar."
            }
        
        # Log the analysis source for debugging
        logger.info(f"Analysis data source: {'MOCK DATA' if using_mock_data else 'REAL AI ANALYSIS'}")
        if analysis_data and not using_mock_data:
            logger.info(f"Analysis summary: {analysis_data.get('summary', 'No summary')[:100]}...")
            
        # Create the response
        response = DocumentResponse(
            document_id=document_id,
            title=f"Belge - {document_id[:8]}",
            document_type=request.template_name,
            content=document_html,  # Include the HTML content directly
            metadata={
                "category": request.case_category,
                "description_length": len(request.case_description),
                "using_mock_data": using_mock_data,
                "real_analysis": not using_mock_data,
                "analysis_timestamp": datetime.now().isoformat()
            },
            analysis=analysis_data
        )
        
        # Special handling for family law/divorce cases: ensure we have meaningful family law content
        family_law_keywords = ['boşan', 'boşanma', 'nafaka', 'velayet', 'çocuk', 'evlilik', 'aile']
        is_family_law_case = any(keyword in request.case_description.lower() for keyword in family_law_keywords)
        
        if is_family_law_case and not using_mock_data:
            # Check if we have family law content in the analysis
            has_family_law_content = False
            
            # Check relevant laws for family law content
            family_law_indicators = ['medeni', 'aile', 'boşanma', 'nafaka', 'velayet']
            
            if analysis_data and "relevant_laws" in analysis_data:
                for law in analysis_data["relevant_laws"]:
                    law_text = str(law.get("title", "")) + str(law.get("description", ""))
                    if any(indicator in law_text.lower() for indicator in family_law_indicators):
                        has_family_law_content = True
                        break
            
            if not has_family_law_content:
                logger.warning("Family law case detected but no family law content in analysis. Adding family law guidance.")
                # Add family law guidance to help the user
                if "relevant_laws" not in analysis_data:
                    analysis_data["relevant_laws"] = []
                
                analysis_data["relevant_laws"].append({
                    "title": "Türk Medeni Kanunu - Aile Hukuku",
                    "description": "Medeni Kanun'un 118-494 maddeleri arasında düzenlenen aile hukuku, evlenme, boşanma, nafaka, velayet gibi aile ilişkilerini düzenler."
                })
                
                # Update the response with the enhanced analysis
                response.analysis = analysis_data
                logger.info("Added family law guidance to the analysis")
        
        return response
    except Exception as e:
        logging.error(f"Error generating document: {str(e)}")
        logging.exception(e)
        raise HTTPException(status_code=500, detail=f"Document generation failed: {str(e)}")

@app.get("/documents/{document_id}/download")
async def download_document(document_id: str):
    """
    Download a generated document
    """
    try:
        # For this simplified version, we'll assume the document is in the output directory
        # with a predictable name pattern
        # In a real system, you'd look up the document path in a database
        
        # Look for the document in the output directory
        output_dir = "app/output"
        for filename in os.listdir(output_dir):
            if document_id in filename:
                file_path = os.path.join(output_dir, filename)
                return FileResponse(
                    path=file_path,
                    filename=filename,
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        
        # If document not found
        raise HTTPException(
            status_code=404,
            detail={"message": "Document not found", "errors": ["The requested document could not be found"]}
        )
        
    except Exception as e:
        logger.error(f"Error downloading document: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail={"message": "Error downloading document", "errors": [str(e)]}
        )

@app.get("/documents/{document_id}/content", response_class=HTMLResponse)
async def get_document_content(document_id: str):
    """
    Get the document content as HTML for direct display
    """
    try:
        # Look for the document in the output directory
        output_dir = "app/output"
        for filename in os.listdir(output_dir):
            if document_id in filename:
                file_path = os.path.join(output_dir, filename)
                # Convert the DOCX to HTML
                html_content = await convert_docx_to_html(file_path)
                return HTMLResponse(content=html_content)
                
        # If document not found
        raise HTTPException(
            status_code=404,
            detail={"message": "Document not found", "errors": ["The requested document could not be found"]}
        )
        
    except Exception as e:
        logger.error(f"Error getting document content: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail={"message": "Error getting document content", "errors": [str(e)]}
        )

async def convert_docx_to_html(file_path):
    """
    Convert a DOCX file to HTML for direct display
    """
    try:
        from docx import Document
        
        # Read the document
        doc = Document(file_path)
        
        # Simple conversion to HTML
        html = ["<div class='document-content'>"]
        
        # Add document paragraphs
        for para in doc.paragraphs:
            # Skip empty paragraphs
            if not para.text.strip():
                html.append("<p>&nbsp;</p>")
                continue
                
            # Check paragraph style and alignment
            alignment = "left"
            if para.alignment == 1:  # WD_ALIGN_PARAGRAPH.CENTER
                alignment = "center"
            elif para.alignment == 2:  # WD_ALIGN_PARAGRAPH.RIGHT
                alignment = "right"
                
            # Start paragraph
            html.append(f"<p style='text-align: {alignment};'>")
            
            # Handle runs with formatting
            for run in para.runs:
                text = run.text.replace('\n', '<br>')
                
                # Skip empty runs
                if not text.strip() and text != "&nbsp;":
                    continue
                    
                # Apply formatting
                if run.bold and run.italic:
                    html.append(f"<strong><em>{text}</em></strong>")
                elif run.bold:
                    html.append(f"<strong>{text}</strong>")
                elif run.italic:
                    html.append(f"<em>{text}</em>")
                else:
                    html.append(text)
                    
            html.append("</p>")
            
        # Add tables
        for table in doc.tables:
            html.append("<table class='table table-bordered'>")
            for row in table.rows:
                html.append("<tr>")
                for cell in row.cells:
                    html.append("<td>")
                    for para in cell.paragraphs:
                        html.append(para.text)
                    html.append("</td>")
                html.append("</tr>")
            html.append("</table>")
            
        html.append("</div>")
        
        return "".join(html)
    except Exception as e:
        logger.error(f"Error converting DOCX to HTML: {str(e)}")
        return f"<div class='error'>Error converting document: {str(e)}</div>"

@app.exception_handler(HTTPException)
async def http_exception_handler(request: Request, exc: HTTPException):
    """
    Custom exception handler for HTTP exceptions
    """
    return JSONResponse(
        status_code=exc.status_code,
        content=exc.detail if isinstance(exc.detail, dict) else {"message": str(exc.detail)}
    )

@app.exception_handler(Exception)
async def general_exception_handler(request: Request, exc: Exception):
    """
    General exception handler
    """
    logger.error(f"Unhandled exception: {str(exc)}")
    return JSONResponse(
        status_code=500,
        content={"message": "An unexpected error occurred", "errors": [str(exc)]}
    )

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000) 